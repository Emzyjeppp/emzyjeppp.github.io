import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_element(name):
    return OxmlElement(name)

def add_heading_styled(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = create_element('w:pBdr')
    bottom = create_element('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_left_right_row(doc, left_text, right_text, left_bold=True, right_bold=True, left_size=10.5, right_size=10):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    
    # Set explicit widths (page width minus margins leaves ~6.9 inches)
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(2.4)
    
    cell_l = table.rows[0].cells[0]
    cell_r = table.rows[0].cells[1]
    
    # Strip cell margins
    for cell in (cell_l, cell_r):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = create_element('w:tcMar')
        for m in ['top', 'left', 'bottom', 'right']:
            node = create_element(f'w:{m}')
            node.set(qn('w:w'), '0')
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)
        
    p_l = cell_l.paragraphs[0]
    p_l.paragraph_format.space_before = Pt(4)
    p_l.paragraph_format.space_after = Pt(2)
    p_l.paragraph_format.keep_with_next = True
    
    run_l = p_l.add_run(left_text)
    run_l.font.name = 'Calibri'
    run_l.font.size = Pt(left_size)
    run_l.bold = left_bold
    
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r.paragraph_format.space_before = Pt(4)
    p_r.paragraph_format.space_after = Pt(2)
    p_r.paragraph_format.keep_with_next = True
    
    run_r = p_r.add_run(right_text)
    run_r.font.name = 'Calibri'
    run_r.font.size = Pt(right_size)
    run_r.bold = right_bold
    return table

def main():
    doc = Document()
    
    # 0.8 inch margins on all sides
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Header Info
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.space_after = Pt(2)
    
    run_name = p_header.add_run("MUHAMMAD JEPRI, A.Md.Kom.\n")
    run_name.bold = True
    run_name.font.name = 'Calibri'
    run_name.font.size = Pt(16)
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(12)
    
    contact_text = (
        "Kotabaru, Kalimantan Selatan  |  082251090558  |  jefryoconner49@gmail.com\n"
        "linkedin.com/in/muhammadjepri  |  github.com/Emzyjeppp  |  jeppp.is-a.dev"
    )
    run_contact = p_contact.add_run(contact_text)
    run_contact.font.name = 'Calibri'
    run_contact.font.size = Pt(9.5)
    
    # Tentang Saya
    add_heading_styled(doc, "TENTANG SAYA")
    p_about = doc.add_paragraph()
    p_about.paragraph_format.space_before = Pt(4)
    p_about.paragraph_format.space_after = Pt(6)
    p_about.paragraph_format.line_spacing = 1.15
    run_about = p_about.add_run(
        "Mahasiswa S-1 Informatika di Universitas Teknologi Digital Indonesia (UTDI) dengan latar belakang D-3 Rekayasa Perangkat Lunak Aplikasi (RPLA). "
        "Memiliki ketertarikan tinggi pada pengembangan backend systems, administrasi data, dan otomatisasi skrip. Memiliki pengalaman kerja praktik "
        "dalam pelayanan administrasi umum, manajemen media informasi publik, serta pengembangan utilitas berbasis Python dan Windows Batch file. "
        "Berorientasi pada detail, adaptif, serta fokus pada penulisan kode yang bersih dan efisien untuk menyelesaikan masalah bisnis."
    )
    run_about.font.name = 'Calibri'
    run_about.font.size = Pt(10)
    
    # Pengalaman Kerja
    add_heading_styled(doc, "PENGALAMAN KERJA")
    
    # Exp 1
    add_left_right_row(doc, "Dinas Ketahanan Pangan dan Pertanian (DKPP) Kotabaru", "31 Januari 2024 - 14 April 2024")
    p_e1_title = doc.add_paragraph()
    p_e1_title.paragraph_format.space_after = Pt(2)
    p_e1_title.paragraph_format.keep_with_next = True
    run_e1_t = p_e1_title.add_run("IT & Multimedia Intern")
    run_e1_t.font.name = 'Calibri'
    run_e1_t.font.size = Pt(10)
    run_e1_t.italic = True
    
    bullets_e1 = [
        "Membantu tugas kepala sub-bagian perencanaan untuk menyusun program kerja dinas serta mengoordinasikan dokumen perencanaan dengan unit kerja lain di lingkungan kantor.",
        "Merancang desain grafis media cetak dan promosi, termasuk pembuatan logo resmi produk benih padi \"Sa'ijaan Unggul\" milik Dinas Ketahanan Pangan dan Pertanian Kotabaru.",
        "Menyunting video kegiatan pertanian dan media komunikasi ucapan hari raya dinas untuk ditayangkan di videotron publik kantor.",
        "Membantu verifikasi serta pengelolaan berkas administrasi umum, invoice pengeluaran, dan surat pertanggungjawaban dinas keuangan dengan akurasi data 100%."
    ]
    for b in bullets_e1:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(2)
        p_b.paragraph_format.left_indent = Inches(0.2)
        run_b = p_b.add_run(b)
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(10)
        
    # Exp 2
    add_left_right_row(doc, "Independent Scripting & Automation Projects", "2023 - Sekarang")
    p_e2_title = doc.add_paragraph()
    p_e2_title.paragraph_format.space_after = Pt(2)
    p_e2_title.paragraph_format.keep_with_next = True
    run_e2_t = p_e2_title.add_run("Self-employed Developer")
    run_e2_t.font.name = 'Calibri'
    run_e2_t.font.size = Pt(10)
    run_e2_t.italic = True
    
    bullets_e2 = [
        "Mengembangkan utilitas desktop berbasis GUI menggunakan Python (CustomTkinter) seperti aplikasi clipboard manager, serta merancang skrip otomatisasi sistem operasi Windows menggunakan Batchfile.",
        "Membangun prototipe dan aplikasi Android native menggunakan Kotlin dan Jetpack Compose (seperti Cupcake-Orders-APP dan pembelajaran Compose dasar) untuk menerapkan arsitektur mobile modern.",
        "Merancang aplikasi web responsif seperti unidaily-asisten-mahasiswa dan portal Pengaduan-Fasilitas-Kampus menggunakan HTML, CSS, PHP, dan MySQL untuk pemenuhan tugas akademik serta pengelolaan basis data relasional."
    ]
    for b in bullets_e2:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(2)
        p_b.paragraph_format.left_indent = Inches(0.2)
        run_b = p_b.add_run(b)
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(10)

    # Exp 3
    add_left_right_row(doc, "Kepolisian Resor (Polres) Kotabaru", "Juli 2018 - Desember 2018 (6 Bulan)")
    p_e3_title = doc.add_paragraph()
    p_e3_title.paragraph_format.space_after = Pt(2)
    p_e3_title.paragraph_format.keep_with_next = True
    run_e3_t = p_e3_title.add_run("Administration & IT Intern (PKL SMK)")
    run_e3_t.font.name = 'Calibri'
    run_e3_t.font.size = Pt(10)
    run_e3_t.italic = True
    
    bullets_e3 = [
        "Membantu Kepala Urusan Umum (Kabag Sium) dalam pelaksanaan tugas-tugas administratif kepolisian secara menyeluruh.",
        "Mencatat rekapitulasi, melakukan penomoran resmi, serta mengarsipkan berkas surat masuk dan surat keluar kantor secara teratur.",
        "Membantu operasional administrasi perkantoran untuk memperlancar komunikasi dokumen dinas di lingkungan Polres Kotabaru."
    ]
    for b in bullets_e3:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(2)
        p_b.paragraph_format.left_indent = Inches(0.2)
        run_b = p_b.add_run(b)
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(10)
        
    # Pendidikan
    add_heading_styled(doc, "PENDIDIKAN")
    
    # Edu 1
    add_left_right_row(doc, "Universitas Teknologi Digital Indonesia (UTDI)", "2025 - Sekarang")
    p_edu1_title = doc.add_paragraph()
    p_edu1_title.paragraph_format.space_after = Pt(2)
    p_edu1_title.paragraph_format.keep_with_next = True
    run_ed1_t = p_edu1_title.add_run("S-1 Informatika - Jalur Rekognisi Pembelajaran Lampau (RPL)")
    run_ed1_t.font.name = 'Calibri'
    run_ed1_t.font.size = Pt(10)
    run_ed1_t.italic = True
    
    p_ed1_b1 = doc.add_paragraph(style='List Bullet')
    p_ed1_b1.paragraph_format.space_after = Pt(2)
    p_ed1_b1.paragraph_format.left_indent = Inches(0.2)
    run_ed1_b1 = p_ed1_b1.add_run("Fokus pada Data Analytics, Machine Learning, Jaringan Syaraf Tiruan, dan Sistem Terdistribusi.")
    run_ed1_b1.font.name = 'Calibri'
    run_ed1_b1.font.size = Pt(10)
    
    p_ed1_b2 = doc.add_paragraph(style='List Bullet')
    p_ed1_b2.paragraph_format.space_after = Pt(2)
    p_ed1_b2.paragraph_format.left_indent = Inches(0.2)
    run_ed1_b2 = p_ed1_b2.add_run("Mata kuliah unggulan: Machine Learning (Nilai A), Data Lakehouse (Nilai A), Jaringan Syaraf Tiruan (Nilai A), Analitika Data (Nilai A).")
    run_ed1_b2.font.name = 'Calibri'
    run_ed1_b2.font.size = Pt(10)
    
    # Edu 2
    add_left_right_row(doc, "Universitas Teknologi Digital Indonesia (UTDI)", "2020 - 2024")
    p_edu2_title = doc.add_paragraph()
    p_edu2_title.paragraph_format.space_after = Pt(2)
    p_edu2_title.paragraph_format.keep_with_next = True
    run_ed2_t = p_edu2_title.add_run("D-3 Rekayasa Perangkat Lunak Aplikasi (RPLA) - A.Md.Kom.")
    run_ed2_t.font.name = 'Calibri'
    run_ed2_t.font.size = Pt(10)
    run_ed2_t.italic = True
    
    bullets_edu2 = [
        "Fokus pada pengembangan aplikasi web client-side & server-side, pemrograman native mobile, basis data non-relasional (NoSQL), serta metodologi pengembangan sistem informasi.",
        "Tugas Akhir: Mengembangkan aplikasi web aspirasi masyarakat terhadap sanggar tari di Yogyakarta menggunakan PHP, MySQL, dan framework Bootstrap (EPrints Repository).",
        "Menyelesaikan proyek aplikasi tim (Proyek Aplikasi 1, 2, & 3) sebagai prasyarat kelulusan akademik."
    ]
    for b in bullets_edu2:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(2)
        p_b.paragraph_format.left_indent = Inches(0.2)
        run_b = p_b.add_run(b)
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(10)

    # Edu 3
    add_left_right_row(doc, "SMK Negeri 1 Kotabaru", "Juli 2017 - Mei 2020")
    p_edu3_title = doc.add_paragraph()
    p_edu3_title.paragraph_format.space_after = Pt(2)
    p_edu3_title.paragraph_format.keep_with_next = True
    run_ed3_t = p_edu3_title.add_run("Teknik Komputer dan Informatika - Rekayasa Perangkat Lunak (RPL)")
    run_ed3_t.font.name = 'Calibri'
    run_ed3_t.font.size = Pt(10)
    run_ed3_t.italic = True
    
    bullets_edu3 = [
        "Mempelajari fondasi teknologi informasi, algoritma dasar pemrograman, perancangan basis data relasional, serta dasar-dasar pemrograman web.",
        "Dinyatakan Lulus pada Mei 2020 dengan penguasaan konsep dasar siklus hidup pengembangan perangkat lunak (SDLC)."
    ]
    for b in bullets_edu3:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(2)
        p_b.paragraph_format.left_indent = Inches(0.2)
        run_b = p_b.add_run(b)
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(10)
        
    # Keterampilan
    add_heading_styled(doc, "KETERAMPILAN & BAHASA")
    
    # 3-column table for aligned colons
    table_sk = doc.add_table(rows=3, cols=3)
    table_sk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_sk.autofit = False
    
    table_sk.columns[0].width = Inches(1.0)
    table_sk.columns[1].width = Inches(0.15)
    table_sk.columns[2].width = Inches(5.75)
    
    sk_data = [
        ("Soft Skill", "Kerja Sama Tim, Pemecahan Masalah (Problem Solving), Kemampuan Adaptasi, Manajemen Waktu, Komunikasi Interpersonal."),
        ("Hard Skill", "Python, CustomTkinter, Batchfile Scripting, HTML/CSS/JavaScript, Bootstrap, Relational Database (SQL), NoSQL Database (MongoDB), Git/GitHub, System Administration."),
        ("Bahasa", "Indonesia (Mahir/Native), Inggris (Menengah).")
    ]
    
    for i, (label, val) in enumerate(sk_data):
        row = table_sk.rows[i]
        cell_lbl = row.cells[0]
        cell_col = row.cells[1]
        cell_val = row.cells[2]
        
        for cell in (cell_lbl, cell_col, cell_val):
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = create_element('w:tcMar')
            for m in ['top', 'left', 'bottom', 'right']:
                node = create_element(f'w:{m}')
                node.set(qn('w:w'), '0')
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)
            
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_before = Pt(2)
        p_lbl.paragraph_format.space_after = Pt(2)
        run_l = p_lbl.add_run(label)
        run_l.font.name = 'Calibri'
        run_l.font.size = Pt(10)
        run_l.bold = True
        
        p_col = cell_col.paragraphs[0]
        p_col.paragraph_format.space_before = Pt(2)
        p_col.paragraph_format.space_after = Pt(2)
        run_c = p_col.add_run(":")
        run_c.font.name = 'Calibri'
        run_c.font.size = Pt(10)
        
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_before = Pt(2)
        p_val.paragraph_format.space_after = Pt(2)
        p_val.paragraph_format.line_spacing = 1.15
        run_v = p_val.add_run(val)
        run_v.font.name = 'Calibri'
        run_v.font.size = Pt(10)
    
    # Sertifikasi
    add_heading_styled(doc, "SERTIFIKASI")
    
    certs = [
        ("Forensik Digital Dasar: Mengamankan Bukti Setelah Insiden Siber — Inixindo Jogja", "April 2026", "ID Sertifikat: 9c1e394a-880d-4562-baed-8b2af54d9beb"),
        ("Prompt Engineering untuk Non-Teknik: Seni Berbicara dengan AI Agar Hasilnya Pas Kebutuhan — Inixindo Jogja", "Maret 2026", "ID Sertifikat: c145ba56-6508-47ae-be13-ad1cd8b27881"),
        ("Implementasi Keamanan OWASP Top 10: Mengidentifikasi dan Mencegah Kerentanan Web Paling Umum — Inixindo Jogja", "Maret 2026", "ID Sertifikat: 631ff6bf-7277-4c1d-84f5-8d8cf48f7737"),
        ("Big Data (Micro Learning) — eduparx (Inixindo Jogja)", "Maret 2026", "ID Kredensial: bbae01e1-85b6-4ec6-ad36-fb1b86a65a88"),
        ("Beyond Coding: Peran Engineer sebagai Arsitek Peradaban di Era AI Otonom — IndoCEISS Jateng", "Februari 2026", "No. Seri: 15/011/INDOCEISS/JATENG/02/2026"),
        ("Deep Learning dalam Pembelajaran Berbasis Teknologi Interactive Flat Panel — ViewSonic Indonesia & HTT-ID", "Desember 2025", "No. Seri: 07.95.003/HTT-ID/12/2025"),
        ("AWS Cloud Campus Class 2025: Unlocking the Power of Cloud — AWS & Metrodata Academy", "Oktober 2025", "Partisipasi Kelas Cloud Campus (Diselenggarakan di UTDI)"),
        ("TOEFL - Like Test — Universitas Teknologi Digital Indonesia (UTDI)", "Januari 2024", "Skor: 440 | No. Seri: 3832/ULDS/I/2024"),
        ("Belajar Dasar Pemrograman Web — Dicoding Indonesia", "Desember 2023", "ID Kredensial: NVP77QY4WPR0")
    ]
    
    for title, date, id_str in certs:
        add_left_right_row(doc, title, date, left_bold=True, right_bold=True, left_size=9.5, right_size=9)
        p_c_id = doc.add_paragraph()
        p_c_id.paragraph_format.space_before = Pt(0)
        p_c_id.paragraph_format.space_after = Pt(2)
        run_cid = p_c_id.add_run(id_str)
        run_cid.italic = True
        run_cid.font.name = 'Calibri'
        run_cid.font.size = Pt(8.5)
        run_cid.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        
    os.makedirs("assets", exist_ok=True)
    doc.save("assets/CV_Muhammad_Jepri.docx")
    print("DOCX successfully generated at assets/CV_Muhammad_Jepri.docx")

if __name__ == "__main__":
    main()
